from docx import Document
import docx.shared
import json

# read config json
with open('config.json') as f:
    config = json.load(f)
    matric_num = config['matric_num']
    module_name = config['module_name']
    sem_num = config['sem_num']
    ay = config['ay']
    assignment_name = config['assignment_name']
    has_appendix = config['has_appendix']
    banned_words = config['banned_words']

# add word count and matric number to last paragraph
def add_word_count(num_words):
    p_text = f"""({num_words} words)
{matric_num}"""
    if not has_appendix:
        p = document.add_paragraph()
        p.text = p_text
    else:
        for p in document.paragraphs:
            if 'Appendix' in p.text:
                p.insert_paragraph_before(p_text)

# open doc
document = Document("essay.docx")

# get styles
style = document.styles['Normal']

# Set line spacing to 2.0
style.paragraph_format.line_spacing = 2.0

# Set font to Times New Roman
style.font.name = 'Times New Roman'

# Set font size to 14
style.font.size = docx.shared.Pt(14)

# Set margins to 1.25"
sections = document.sections
for section in sections:
    section.top_margin = docx.shared.Inches(1.25)
    section.bottom_margin = docx.shared.Inches(1.25)
    section.left_margin = docx.shared.Inches(1.25)
    section.right_margin = docx.shared.Inches(1.25)

# Add header
header_text = f"""{module_name}
Sem {sem_num} AY {ay}
{assignment_name}
"""
header = document.sections[0].header
p = header.paragraphs[0]
p.text = header_text

# set header style
style = document.styles['Header']
style.font.name = 'Times New Roman'
style.font.size = docx.shared.Pt(12)

# get all text as string
text = []
for p in document.paragraphs:
    if has_appendix and 'Appendix' in p.text:
        break
    text.append(p.text)

# add word count and matric number
flat_text = ' '.join(text)
num_words = len(flat_text.split())
add_word_count(num_words)

# TODO: Add page numbers

# check for banned words
if any(word in flat_text for word in banned_words):
    print("You have used banned words. Please remove them and try again.")

# TODO: Check for free-floating pointer words

document.save("formatted-essay.docx")